import os
import re
import pandas as pd
# do not show annoyng subset warnings
pd.options.mode.chained_assignment = None  # default='warn'
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.lines as lines
from matplotlib.ticker import MultipleLocator
from matplotlib.ticker import FuncFormatter
from matplotlib import rcParams
import datetime
import lib.plotformat as pformat
import lib.doc29lib as ld


from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from scipy.misc import imread
from scipy.interpolate import RectBivariateSpline
#from geopandas import GeoDataFrame
from descartes import PolygonPatch

def DENverdeling_evaluatie(df_in,hdr_time,hdr_sum,hdr_LT,hdr_multiindex):

    #%% make sure changes in this function do not affect df
    df = df_in
    #%% now perform operations
    
    hdr1 = hdr_multiindex+ ', landingen'
    hdr2 = hdr_multiindex+ ', starts'
    hdr3 = hdr_multiindex+ ', totaal'
    
    
    df[hdr_time] = pd.to_datetime(df[hdr_time])
    df['DEN'] =df[hdr_time]
    df['DEN'] = "D"
    # change string to date
    t1 = pd.to_datetime('22:59:59')
    t2 = pd.to_datetime('07:00:00')
    t3 = pd.to_datetime('18:59:59')
    t4 = pd.to_datetime('23:00:00')
    t5 = pd.to_datetime('05:59:59')
    # logical indexing
    df['DEN'][(df[hdr_time]>t1) | (df[hdr_time]<t2)] = "N"
    df['DEN'][(df[hdr_time]>t3) & (df[hdr_time]<t4)] = "E"
    df['DEN'][(df[hdr_time]>t5) & (df[hdr_time]<t2)] = "EM"
    # change LT or AD to starts en landingen
    df[hdr_LT][(df[hdr_LT] == 'L') | (df[hdr_LT] == 'A')] = hdr1
    df[hdr_LT][(df[hdr_LT] == 'T') | (df[hdr_LT] == 'D')] =  hdr2
    # use groubpy to aggregate
    df_out          = pd.pivot_table(df,values=hdr_sum,index='DEN',columns=hdr_LT, aggfunc=np.sum) 
    # swap EM and N rows
    df_out = df_out.reindex(["D", "E", "N","EM"])
    
    
    df_out[hdr3]    = df_out[hdr1]+df_out[hdr2]
    # swap rows
    
    df_out.loc['totaal']    = df_out.sum()
    # round to 100 
    df_out = round(df_out,-2)
    
    #create group headers
    a = df_out .columns.str.split(', ', expand=True).values
    df_out.columns = pd.MultiIndex.from_tuples([('', x[0]) if pd.isnull(x[1]) else x for x in a])
    return df_out


def DENverdeling(df_in,
                 hdr_time,
                 hdr_LT,
                 hdr_multiindex,
                 hdr_sum=None):

    #%% make sure changes in this function do not affect df
    df = df_in
    #%% now perform operations
    
    hdr1 = hdr_multiindex+ ', landingen'
    hdr2 = hdr_multiindex+ ', starts'
    hdr3 = hdr_multiindex+ ', totaal'
    
    df[hdr_time]    = pd.to_datetime(df[hdr_time])
    df['DEN']       =   df[hdr_time]
    df['DEN']       = "D"
    
    # logical indexing
    df['DEN'][(df[hdr_time].dt.hour>22) | (df[hdr_time].dt.hour<7)] = "N"
    df['DEN'][(df[hdr_time].dt.hour>18) & (df[hdr_time].dt.hour<23)] = "E"
    df['DEN'][(df[hdr_time].dt.hour>5) & (df[hdr_time].dt.hour<7)] = "EM"
    
    # change LT or AD to starts en landingen
    df[hdr_LT][(df[hdr_LT] == 'L') | (df[hdr_LT] == 'A')] = hdr1
    df[hdr_LT][(df[hdr_LT] == 'T') | (df[hdr_LT] == 'D')] =  hdr2

    # use groubpy to aggregate
    if hdr_sum:
        df_out          = pd.pivot_table(df,values=hdr_sum,index='DEN',columns=hdr_LT, aggfunc=np.sum) 
    else:
        df_out          = pd.pivot_table(df,values=hdr_time,index='DEN',columns=hdr_LT, aggfunc='count') 
    
    # swap EM and N rows
    df_out = df_out.reindex(["D", "E", "N","EM"])
    # total columns
    df_out[hdr3] = df_out[hdr1]+df_out[hdr2]
    
    # total row
    total = df_out.sum(numeric_only=True)
    df_out = df_out.append(total, ignore_index=True)
    
    #create group headers
    a = df_out .columns.str.split(', ', expand=True).values
    df_out.columns = pd.MultiIndex.from_tuples([('', x[0]) if pd.isnull(x[1]) else x for x in a])
    
    # round
    df_out = df_out.round()
    df_out = df_out.set_index([pd.Index(["dag 07-19 uur", 
                                         "avond 19-23 uur", 
                                         "nacht 23-06 uur",
                                         "vroege ochtend 06-07 uur",
                                         "totaal"])])
    # round to 100 
    df_out = round(df_out,-2)

    return df_out



def SWverdeling(df_in,hdr_date,gj):
    #%% make sure changes in this function do not affect df
    df = df_in
    
    #%% now perform operations
    #make summer/winter column
    df[hdr_date] = pd.to_datetime(df[hdr_date])
    df['SW']  =df[hdr_date]
    df['SW'] = "zomer"
    
    season = 'zomer'
    start = gebruiksjaar(gj,season)
    df['SW'][df[hdr_date]<start] = 'winter'
    df_out = df.groupby(['SW']).agg('count')
    df_out = df_out[hdr_date]
    
    # round to 100 
    df_out = round(df_out,-2)

    return df_out


def gebruiksjaar(year,season):
    if (season =='summer') | (season == 'zomer'):
        #zondag van het laatste weekend van Maart
        date = datetime.date(year,  3, 31 - (int(float(4 + 5 * year/4)) % 7))
    else:
        #zondag van het laatste weekend van oktober
        date = datetime.date(year, 10, 31 - (int(float(1 + 5 * year/4)) % 7)) 
    startdate = datetime.datetime.combine(date,datetime.time(0,0))
    return startdate
    

def plot_baangebruik_evaluatie(trf_files,
                     labels,
                     TL,
                     trf_realisatie=None,
                     DEN='DEN',
                     fname=None,
                     n=7,
                     runways=None,
                     ylim=[0,110000],
                     dy=10000,
                     reftraffic=1,
                     style='MER',
                     dpi=300):
    
    def NumberFormatter(x, pos):
        'The two args are the value and tick position'
        return '{:,.0f}'.format(x).replace(',', '.')
    def GetVal(var, i):
        'Scalar of een list'
        if isinstance(var, list):
            i = min(i, len(var)-1) # hergebruik van de laatste waarde 
            return var[i]
        else:
            return var

    
    # check of trf_files een string of list is   
    if not isinstance(trf_files, list):
        trf_files = [trf_files]
        
    # N = trf_stats['d_lt'].count()
    
    # X-positie van de bars
    x = np.arange(n)

    ntrf = len(trf_files)
    i = ntrf - 1
    w = GetVal(pformat.baangebruik[style]['barwidth'], i) # of /ntrf
    g = GetVal(pformat.baangebruik[style]['bargap'], i) # of /ntrf?
    
    dx = [(w+g)*(i - 0.5*(ntrf-1)) for i in range(ntrf)]
    
    # markers en staafjes
    marker_height = dy * GetVal(pformat.baangebruik[style]['markerheight'], i)
    mw = GetVal(pformat.baangebruik[style]['markerwidth'], i)
    dxm = list(dx)
    
    # clip marker
    if ntrf == 2:
        mw = (mw + w)/2
        dxm[0] = dx[0] - (mw-w)/2
        dxm[1] = dx[1] + (mw-w)/2
    elif ntrf > 2:
        mw = [min(mw, w+g)]
    
    # open 1 figuur
    fig, (ax1) = plt.subplots(1, 1, sharey=True)
    fig.set_size_inches(21/2.54, 10/2.54)
    
    # margins
    fig.subplots_adjust(bottom=0.18)    
    fig.subplots_adjust(wspace=0)
    
    # LEGENDA 1
    ax0 = fig.add_axes([0, 0.9, 0.6, 0.1]) 
    # geen assen
    ax0.axis('off') 
    # genormaliseerde asses
    ax0.set_xlim(-.5, 0.5)
    ax0.set_ylim(0, 1)
    
    
    for i, yi, xt, stylo1, stylo2 in [(0, 0.5, 0,'bar','marker'),
                       (1, 0.5, 0.15,'refbar','refmarker'),
                       (2, 0.5, 0.25,'refbar1','refmarker1'),
                       (3, 0.5, 0.45,'refbar2','refmarker2')]:
        if i<len(labels):
            if i<2:
                ax0.bar(xt, height=0.6, bottom=0.1,
                        width=0.02,
                        **pformat.baangebruik[style][stylo1],
                        zorder=4)
            
            ax0.bar(xt, height=0.1, bottom=0.5,
                    width=0.04,
                    **pformat.baangebruik[style][stylo2],
                    zorder=6)
            
            ax0.text(xt+0.03, 0.5, labels[i],
                     **pformat.baangebruik[style]['legendtext'])
    
    
    if trf_realisatie:
        # LEGENDA 2 - REALISATIE
        ax0b = fig.add_axes([-0.05, 0.9, 0.3, 0.1]) 
        # geen assen
        ax0b.axis('off')
        # genormaliseerde asses
        ax0b.set_xlim(-.5, 0.5)
        ax0b.set_ylim(0, 1)
        #plot realisatie
        ax0b.plot([0,0.3],
                 [0.5,0.5],
                 'k--',
                 zorder=10,
                 linewidth=2)   
        
        ax0b.text(0.3, 0.5, 'realisatie',
                     **pformat.baangebruik[style]['legendtext'])
    
    
        # Realisatie
        trf = pd.read_csv(trf_realisatie, delimiter='\t')
        
        # if night figure do
        if DEN!='DEN':
            trf = trf[trf['d_den']=='N']
              
        trf = trf.groupby(['d_lt', 'd_runway', 'd_myear'])['total'].sum().reset_index()
    
        trf_stats = trf.groupby(['d_lt', 'd_runway'])['total'].agg(['min','max','mean']).reset_index()
    #    print(trf_stats)
    
        # selecteer L of T
        trfn = trf_stats.loc[trf_stats['d_lt'] == TL]
        trfn = trfn.sort_values(by =['mean'],ascending=False)
        trfn = trfn.head(n) # gaat alleen goed als er ook echt n-runways zijn
        ind= trfn['d_runway']
    #    print(trfn)
    #    print(ind)
   
    stylo1 = ['bar','refbar','refbar1','refbar2']
    stylo2 = ['marker','refmarker','refmarker1','refmarker2']
    
    # verwerken traffics
    for i, trf_file in enumerate(trf_files):
        # lees csv
        trf = pd.read_csv(trf_file, delimiter='\t') 
        
        # if night figure do
        if DEN!='DEN':
            trf = trf[trf['d_den']=='N']
        
        # aggregeer etmaalperiode en bereken stats
        trf = trf.groupby(['d_lt', 'd_runway', 'd_myear'])['total'].sum().reset_index()
        trf_stats = trf.groupby(['d_lt', 'd_runway'])['total'].agg(['min','max','mean']).reset_index()
        
        
        # filter out NOG flights
        trf_stats =trf_stats.query("d_runway != 'NOG'").reset_index()
        
        
        # sorteer
        if 'key' not in trf_stats.columns:
            trf_stats['key'] = trf_stats['d_lt'] + trf_stats['d_runway']

        if runways is not None:
            # tweede traffic in dezelfde volgorde
            keys = [k + r for k in runways for r in runways[k]]    # keys: combinatie van lt en runway
            sorterIndex = dict(zip(keys, range(len(keys))))        # plak een volgnummer aan de keys     
            trf_stats['order'] = trf_stats['key'].map(sorterIndex) # soteerindex toevoegen
            trf_stats = trf_stats.sort_values(by=['order'])        # sorteer dataframe
        else:
            trf_stats = trf_stats.sort_values(by=['d_lt', 'mean'], ascending=False)
            runways = {'L': trf_stats['d_runway'].loc[trf_stats['d_lt'] == 'L'],
                       'T': trf_stats['d_runway'].loc[trf_stats['d_lt'] == 'T']}
        
        if TL == 'L':
            le = ['landingen']
            st1 = 'facecolor2'
        else:
            le = ['starts']
            st1 ='facecolor1'
            
        # maak de plot
        for lt, xlabel, fc, spine, ax in zip(TL,
                                             le,
                                             pformat.baangebruik[style][st1],
                                             ['right', 'left'],
                                             [ax1]):

            # selecteer L of T
            trf2 = trf_stats.loc[trf_stats['d_lt'] == lt]
            trf2 = trf2.head(n) # gaat alleen goed als er ook echt n-runways zijn
#            print(trf2)
            trf2 = trf2.set_index('d_runway')
            #set index to correct order
#            trf2 = trf2.reindex(ind)
            trf2 = trf2.reset_index()
            
#            print(trf2)
            
            # staafjes
            bar_height = trf2['max'] - trf2['min']
                
            ax.bar(x+dx[i], height=bar_height, bottom=trf2['min'],
                   width=w,
                   **pformat.baangebruik[style][stylo1[i]],
                   zorder=4)
            
            # gemiddelde
            ax.bar(x+dxm[i], height=marker_height, bottom=trf2['mean']-marker_height/2,
                   width=mw,
                   **pformat.baangebruik[style][stylo2[i]],
                   zorder=4)
            
          
            # achtergrondkleur
            ax.set_facecolor(fc)
        
            # border
            plt.setp(ax.spines.values(), **pformat.baangebruik[style]['spines'])
                     
            # scheidingslijntje tussen subplots
            ax.spines[spine].set_color('white')
        
            # tweak scheidingslijntje tussen subplots
            for y in [0,1]:
                frame = lines.Line2D([0, 1], [y, y], 
                              transform=ax1.transAxes,
                              **pformat.baangebruik[style]['spines'],
                              zorder=10)
                frame.set_clip_on(False)
                ax.add_line(frame)  
                 
            # gridlines
            ax.grid(which='major', axis='y', linewidth=0.5, color='white')
            
            # geen tickmarks
            plt.setp([ax.get_xticklines(), ax.get_yticklines()], color='none')
        
            # label size and color   
            ax.tick_params(axis='both', **pformat.baangebruik[style]['axislabel'])
                    
            # X-as
            ax.set_xticks(x)
            ax.set_xticklabels(trf2['d_runway'])
            ax.text(0.5, -0.18, xlabel,
               transform=ax.transAxes,
               **pformat.baangebruik[style]['grouptext'])
          
            # X-as lijntjes
            ax.set_xlim(-0.5, n-0.5)
            line = lines.Line2D([0.02, 0.98], [-.11, -.11], 
                                transform=ax.transAxes,
                                **pformat.baangebruik[style]['grouplines'])
            line.set_clip_on(False)
            ax.add_line(line)
        
            # Y-as
            ax.set_ylim(ylim)
            ax.yaxis.set_major_locator(MultipleLocator(base=dy))
            ax.yaxis.set_major_formatter(FuncFormatter(NumberFormatter))
      
    if trf_realisatie:
        #plot realisatie
        ax1.plot([x-0.4,x+0.4],
                 [trfn['mean'],trfn['mean']],
                 'k--',
                 zorder=10,
                 linewidth=2)        
    
    #save if nescessary, or print to screen
    if fname:
        fig.savefig(fname, dpi=dpi)
        plt.close(fig)
    else:
        plt.show()

def df2table(t,df,style_header,style_row1,style_body):
    headers =  df.columns.values
    # print headers
    for c,header in zip(t.rows[0].cells,headers):
        c.paragraphs[0].add_run(str(header))
        c.paragraphs[0].style = style_header
    i=0
    
    # print body
    for row in t.rows:
        j=0
        if i > 0:
            for cell in row.cells:
                header = headers[j]
                if j==0:
                    if type(df[header][i-1]) is float or type(df[header][i-1]) is np.float64:
                        cell.paragraphs[0].add_run("{0:.0f}".format(round(df[header][i-1],-2)))
                    else:
                        cell.paragraphs[0].add_run(str(df[header][i-1]))    
                    
                    cell.paragraphs[0].style = style_row1
                else:
                    if type(df[header][i-1]) is float or type(df[header][i-1]) is np.float64:
                        cell.paragraphs[0].add_run("{0:.0f}".format(round(df[header][i-1],-2)))
                    else:
                        cell.paragraphs[0].add_run(str(df[header][i-1])) 
                        
                    cell.paragraphs[0].style = style_body
    
                j+=1               
        i+=1  


def fig22(DEN_zw,fn,hs):
    c1 = hs['kleuren']['schipholblauw']
    c2 = hs['kleuren']['middagblauw']
    c3 = hs['kleuren']['wolkengrijs_1']
    c4 = hs['kleuren']['schemergroen']
    
    font = hs['fontname']['grafiek']
    fontsize = hs['size']['font']
    
    r =range(len(DEN_zw))
    names =['landingen','starts','landingen','starts']
    
    fig1, ax1 = plt.subplots(figsize=(12, 4))
    
    ax1.barh(r,
         DEN_zw['D'], 
         align='center',
         color=c1,
         label = 'dag')
    left = DEN_zw['D']
    ax1.barh(r,
         DEN_zw['E'],  
         left = left,
         align='center',
         color=c2,
         label = 'avond')
    left = left + DEN_zw['E']
    ax1.barh(r,
         DEN_zw['N'], 
         left = left,
         align='center',
         color=c3,
         label = 'nacht')
    left = left + DEN_zw['N']
    ax1.barh(r,
         DEN_zw['EM'],
         left = left,
         align='center',
         color=c4,
         label = 'vroege ochtend')
    
    plt.yticks(r, 
               names,              
               fontname = font,
               fontsize = fontsize)
    ax1.grid(axis='x')
    ax1.set_ylabel('Zomer              Winter',
                  fontname = font,
                  fontsize = fontsize)
    
    
    # get rid of box around grpah
    for spine in plt.gca().spines.values():
        spine.set_visible(False) #Indentation updated..
    # Add a legend
    plt.legend(prop={'family':font, 'size':fontsize},
               ncol=4,
               bbox_to_anchor=(0.9, 1.15))
    # save
    plt.savefig(fn,dpi=300, bbox_inches='tight')
    
    plt.show()
    
#def fig23(DEN_zw,fn,hs):
#    c1 = hs['kleuren']['schipholblauw']
#    c2 = hs['kleuren']['middagblauw']
#    c3 = hs['kleuren']['wolkengrijs_1']
#    c4 = hs['kleuren']['schemergroen']
#    
#    font = hs['fontname']['grafiek']
#    fontsize = hs['size']['font']
#    
#    r =range(len(DEN_zw))
#    names =DEN_zw.index.tolist()
#    
#    fig1, ax1 = plt.subplots(figsize=(12, 4))
#    ax1.grid(axis='y')
#   
#    
#    ax1.bar(r,
#         DEN_zw, 
#         align='center',
#         color=c1,
#         width= 0.5)
#
#    plt.xticks(r, 
#               names,              
#               fontname = font,
#               fontsize = fontsize)
#    
#    vals = ax1.get_yticks()
#    ax1.set_yticklabels(['{0:.0f}%'.format(x) for x in vals],              
#               fontname = font,
#               fontsize = fontsize)
#     
#    plt.xlabel('Maximum startgewicht in tonnen',               
#               fontname = font,
#               fontsize = fontsize)
# 
#    # get rid of box around grpah
#    for spine in plt.gca().spines.values():
#        spine.set_visible(False) #Indentation updated..
#
#    # save
#    plt.savefig(fn,dpi=300, bbox_inches='tight')
#    
#    plt.show()
    
def tab42(tf_pref,baancombinaties):
    tf_pref.loc[tf_pref['d_schedule']==6,'d_den']='EM'
    
    # nacht
    tf_pref_n = tf_pref[tf_pref['d_den']=='N']
    tf_pref_n = pd.merge(tf_pref_n,baancombinaties,left_on=['d_combination','d_den'],right_on=['combination','period'],how='left')
    tf_pref_n = tf_pref_n.fillna('-')
    bcverdeling_nacht = tf_pref_n.groupby(['preference'])['total'].sum()
    bc_n_relatief = bcverdeling_nacht/bcverdeling_nacht.sum()*100
    
    # make clean dataframe
    d = {'Preferentie':[1,2,3,4,'Subtotaal','Anders','Totaal'],
         'Inzet': [bcverdeling_nacht[1],
                   bcverdeling_nacht[2],
                   bcverdeling_nacht[3],
                   bcverdeling_nacht[4],
                   bcverdeling_nacht[1:5].sum(),
                   bcverdeling_nacht[0],
                   bcverdeling_nacht.sum()],
         'Relatief [%]': ['{0:.1f}%'.format(bc_n_relatief[1]),
                   '{0:.1f}%'.format(bc_n_relatief[2]),
                   '{0:.1f}%'.format(bc_n_relatief[3]),
                   '{0:.1f}%'.format(bc_n_relatief[4]),
                   '{0:.1f}%'.format(bc_n_relatief[1:5].sum()),
                   '{0:.1f}%'.format(bc_n_relatief[0]),
                   '{0:.1f}%'.format(bc_n_relatief.sum())]}
    tabel42b = pd.DataFrame(data=d)
    tabel42b = tabel42b[['Preferentie', 'Inzet', 'Relatief [%]']]
    
    
    # dag verdeling 
    tf_pref = tf_pref[(tf_pref['d_den']=='D') | (tf_pref['d_den']=='E') | (tf_pref['d_den']=='EM')]
    baancombinaties = baancombinaties[baancombinaties['period']!='N']
    tf_pref = pd.merge(tf_pref,baancombinaties,left_on=['d_combination'],right_on=['combination'],how='left')
    tf_pref = tf_pref.fillna('-')
    bcverdeling = tf_pref.groupby(['preference'])['total'].sum()
    bc_relatief = bcverdeling/bcverdeling.sum()*100
    
    # make clean dataframe
    d = {'Preferentie':[1,2,3,4,5,6,'Subtotaal','Anders','Totaal'],
         'Inzet': [bcverdeling[1],
                   bcverdeling[2],
                   bcverdeling[3],
                   bcverdeling[4],
                   bcverdeling[5],
                   bcverdeling[6],
                   bcverdeling[1:7].sum(),
                   bcverdeling[0],
                   bcverdeling.sum()],
         'Relatief [%]': ['{0:.1f}%'.format(bc_relatief[1]),
                   '{0:.1f}%'.format(bc_relatief[2]),
                   '{0:.1f}%'.format(bc_relatief[3]),
                   '{0:.1f}%'.format(bc_relatief[4]),
                   '{0:.1f}%'.format(bc_relatief[5]),
                   '{0:.1f}%'.format(bc_relatief[6]),
                   '{0:.1f}%'.format(bc_relatief[1:7].sum()),
                   '{0:.1f}%'.format(bc_relatief[0]),
                   '{0:.1f}%'.format(bc_relatief.sum())]}
         
    tabel42a = pd.DataFrame(data=d)
    tabel42a = tabel42a[['Preferentie', 'Inzet', 'Relatief [%]']]
    
    return tabel42a, tabel42b

def plot_baangebruik(trf_files,
                     labels,
                     huisstijl,
                     trf_realisatie = None,
                     den=['D', 'E', 'N'],
                     fname=None,
                     n=7,
                     runways=None,
                     ylim=[0,110000],
                     dy=10000,
                     reftraffic=1,
                     style='MER',
                     dpi=300):
    '''Plot het baangebruik'''
    
    c1 = huisstijl['kleuren']['schipholblauw']
    c2 = huisstijl['kleuren']['middagblauw']
    c3 = huisstijl['kleuren']['wolkengrijs_1']
    c4 = huisstijl['kleuren']['schemergroen']
    
    font = huisstijl['fontname']['grafiek']
    fontsize = huisstijl['size']['font']
    
    
    def NumberFormatter(x, pos):
        'The two args are the value and tick position'
        return '{:,.0f}'.format(x).replace(',', '.')
    def GetVal(var, i):
        'Scalar of een list'
        if isinstance(var, list):
            i = min(i, len(var)-1) # hergebruik van de laatste waarde 
            return var[i]
        else:
            return var

    
    # check of trf_files een string of list is   
    if not isinstance(trf_files, list):
        trf_files = [trf_files]
        
    # N = trf_stats['d_lt'].count()
    
    # X-positie van de bars
    x = np.arange(n)

    ntrf = len(trf_files)
    i = ntrf - 1
    w = (GetVal(pformat.baangebruik[style]['barwidth'], i) # of /ntrf
         * n/7) # normaliseer voor de aslengte
    g = GetVal(pformat.baangebruik[style]['bargap'], i) # of /ntrf?
    
    dx = [(w+g)*(i - 0.5*(ntrf-1)) for i in range(ntrf)]
    
    # markers en staafjes
    marker_height = (GetVal(pformat.baangebruik[style]['markerheight'], i)
                     * (ylim[-1] - ylim[0]) / 10) 
    mw = (GetVal(pformat.baangebruik[style]['markerwidth'], i)
          * n/7)
    dxm = list(dx)
    
    # clip marker
    if ntrf == 2:
        mw = (mw + w)/2
        dxm[0] = dx[0] - (mw-w)/2
        dxm[1] = dx[1] + (mw-w)/2
    elif ntrf > 2:
        mw = [min(mw, w+g)]
    
    # twee aansluitende subplots
    fig, (ax1, ax2) = plt.subplots(1, 2, sharey=True)
    fig.set_size_inches(12, 6)
    
    # margins
    fig.subplots_adjust(bottom=0.18)    
    fig.subplots_adjust(wspace=0)
    
    # legenda
    ax0 = fig.add_axes([0.7, 0.9, 0.05, 0.1]) 
    
    # geen assen
    ax0.axis('off')
    
    # genormaliseerde asses
    ax0.set_xlim(-0.5*n/7, 0.5*n/7)
    ax0.set_ylim(0, 1)
    
    # staafjes
    if ntrf == 2:
        #TODO: 1 of >2 staafjes
        # gemiddelde
        for i, yi, bottom, xt, yt, alignment in [(0, 0.4, 0.1, 0.2, 0.4, 'right'),
                                                 (1, 0.5, 0.3, 0.8, 0.4, 'left')]:
            if i == reftraffic:
                c_bar=c2
                c_marker=c1
            else:
                c_bar=c3
                c_marker=c4
            ax0.bar(dx[i], height=0.6, bottom=bottom,
                    width=w,
                    edgecolor=c_bar,
                    facecolor=c_bar,
                    linewidth=0.3,
                    zorder=4)
            ax0.bar(dxm[i], height=0.05, bottom=yi,
                    width=mw,
                    edgecolor=c_marker,
                    facecolor=c_marker,
                    linewidth=0.3,
                    zorder=6)
            ax0.text(xt, yt, labels[i],
               transform=ax0.transAxes,
               horizontalalignment=alignment,
               fontsize=fontsize,
               verticalalignment='center',
               fontname = font)
    
    # LEGENDA 2 - REALISATIE
    ax0b = fig.add_axes([-0.05, 0.9, 0.3, 0.1]) 
    # geen assen
    ax0b.axis('off')
    # genormaliseerde asses
    ax0b.set_xlim(-.5, 0.5)
    ax0b.set_ylim(0, 1)
    #plot realisatie
    ax0b.plot([0,0.2],
             [0.5,0.5],
             'k--',
             zorder=10,
             linewidth=2)   
    
    ax0b.text(0.7, 0.5, 'Realisatie',
              transform=ax0b.transAxes,
              horizontalalignment='left',
              fontsize=fontsize,
              verticalalignment='center',
              fontname = font)
    
    #realisatie
    if den == ['N']:
        trf_realisatie = trf_realisatie.loc[(trf_realisatie['DEN']=='N') | (trf_realisatie['DEN']=='EM'),:]
    
    
    trf_realisatie = trf_realisatie.groupby(['C_AD','C_runway'])['C_actual'].count().reset_index()
    trf_realisatie.loc[trf_realisatie['C_AD']=='realisatie, landingen','C_AD'] ='L'
    trf_realisatie.loc[trf_realisatie['C_AD']=='realisatie, starts','C_AD'] ='T'
    
    # sorteer
    if 'key' not in trf_realisatie.columns:
        trf_realisatie['key'] = trf_realisatie['C_AD'] + trf_realisatie['C_runway']


    
    # drop runways that are not included in prognosis
    trf_realisatie=trf_realisatie.dropna()
    
    # verwerken traffics
    for i, trf_file in enumerate(trf_files):
        
        # lees csv
        trf = pd.read_csv(trf_file, delimiter='\t')
        trf = trf.loc[trf['d_den'].isin(den)]
        
        # aggregeer etmaalperiode en bereken stats
        trf = trf.groupby(['d_lt', 'd_runway', 'd_myear'])['total'].sum().reset_index()
        trf_stats = trf.groupby(['d_lt', 'd_runway'])['total'].agg(['min','max','mean']).reset_index()
        
        # sorteer
        if 'key' not in trf_stats.columns:
            trf_stats['key'] = trf_stats['d_lt'] + trf_stats['d_runway']

        if runways is not None:
            # tweede traffic in dezelfde volgorde
            keys = [k + r for k in runways for r in runways[k]]    # keys: combinatie van lt en runway
            sorterIndex = dict(zip(keys, range(len(keys))))        # plak een volgnummer aan de keys     
            trf_stats['order'] = trf_stats['key'].map(sorterIndex) # soteerindex toevoegen
            trf_stats = trf_stats.sort_values(by=['order'])        # sorteer dataframe
        else:
            trf_stats = trf_stats.sort_values(by=['d_lt', 'mean'], ascending=False)
            runways = {'L': trf_stats['d_runway'].loc[trf_stats['d_lt'] == 'L'],
                       'T': trf_stats['d_runway'].loc[trf_stats['d_lt'] == 'T']}
        
        # maak de plot
        for lt, xlabel, fc, spine, ax in zip(['T', 'L'],
                                             ['starts', 'landingen'],
                                             [c1,c2],
                                             ['right', 'left'],
                                             [ax1, ax2]):
            
            # selecteer L of T
            trf2 = trf_stats.loc[trf_stats['d_lt'] == lt]
            trf2 = trf2.head(n) # gaat alleen goed als er ook echt n-runways zijn
            
            # staafjes
            bar_height = trf2['max'] - trf2['min']
            if i == reftraffic:
                c_bar=c2
                c_marker=c1
            else:
                c_bar=c3
                c_marker=c4
                
            ax.bar(x+dx[i], height=bar_height, bottom=trf2['min'],
                   width=w,
                   edgecolor=c_bar,
                   facecolor=c_bar,
                   linewidth=0.3,
                   zorder=4)
            
            # gemiddelde
            ax.bar(x+dxm[i], height=marker_height, bottom=trf2['mean']-marker_height/2,
                   width=mw,
                   edgecolor=c_marker,
                   facecolor=c_marker,
                   linewidth=0.3,
                   zorder=4)
            

            # TO DO: volgorde van de realisatie NIET altijd hetzelfde als de prognose
            trf_r = trf_realisatie.loc[trf_realisatie['C_AD']==lt,:].nlargest(n,'C_actual')
#            print(trf_r)
            ax.plot([x-0.4,x+0.4],
                    [trf_r['C_actual'],trf_r['C_actual']],
                    'k--',
                    zorder=10,
                    linewidth=2)  
            
            
            # get rid of box around grpah
            for sp in ax.spines.values():
                sp.set_visible(False) #Indentation updated..
                 
#            # geen tickmarks
            plt.setp([ax.get_xticklines(), ax.get_yticklines()], color='none')
            
            
            # scheidingslijntje tussen subplots
            ax.spines[spine].set_visible(True)
            ax.spines[spine].set_color('grey')
            
            # gridlines
            ax.grid(which='major', axis='y', linewidth=0.5)
            
            # label size and color   
            ax.tick_params(axis='both',
                           size=fontsize)
                    
            # X-as
            ax.set_xticks(x)
            ax.set_xticklabels(trf2['d_runway'])
            ax.text(0.5, -0.18, xlabel,
               transform=ax.transAxes,
               fontname=font,
               fontsize=fontsize)
          
            # X-as lijntjes
            ax.set_xlim(-0.5, n-0.5)
            line = lines.Line2D([0.02, 0.98], [-.13, -.13], 
                                transform=ax.transAxes,
                                **pformat.baangebruik[style]['grouplines'])
            line.set_clip_on(False)
            ax.add_line(line)
        
            # Y-as
            ax.set_ylim(ylim)
            ax.yaxis.set_major_locator(MultipleLocator(base=dy))
            ax.yaxis.set_major_formatter(FuncFormatter(NumberFormatter))
    


        
    if fname:
        
        fig.savefig(fname, dpi=dpi, bbox_inches='tight')
        plt.show()
        plt.close(fig)
    else:
        plt.show()
        
def appendDF(fns):
    for i, fn in enumerate(fns):
        if i == 0:    
            df = pd.read_csv(fn, delim_whitespace = 1)
        else:
            df = df.append(pd.read_csv(fn, delim_whitespace = 1))
    return df



def plotContour(ax,hdr,dat,dB,c1,c2):
    for d in dat['dat']:
        X, Y, Z = ld.verfijn(hdr, d, k=20)
        cs = ax.contour(X, Y, Z, 
                        levels=[dB], 
                        colors=c2, 
                        linewidths=[3, 3],
                        alpha=0.1)

    X1, Y1, Z1 = ld.verfijn(hdr, dat['mean'], k=20)
    X2, Y2, Z2 = ld.verfijn(hdr, dat['dhi'], k=20)
    X3, Y3, Z3 = ld.verfijn(hdr, dat['dlo'], k=20)
    #
    cs1 = ax.contour(X1, Y1, Z1, levels=[dB], colors=c1, linewidths=[1, 1])
    cs2 = ax.contour(X2, Y2, Z2, levels=[dB], colors=c2, linewidths=[0.5, 0.5])
    cs3 = ax.contour(X3, Y3, Z3, levels=[dB], colors=c2, linewidths=[0.5, 0.5])
    
    return cs1
    
def figHistory(history,prog,key,fn,hs,ylim=None):
        
    #  unpack hs    
    font = hs['fontname']['grafiek']
    fontsize = hs['size']['font']
    linesize = hs['size']['line']
    markersize = hs['size']['marker']
    
    c1 = hs['kleuren']['schipholblauw']
    c2 = hs['kleuren']['middagblauw']
    c3 = hs['kleuren']['wolkengrijs_1']
    c4 = hs['kleuren']['schemergroen']
    
    #  unpack history
    realisatie = history.loc[:, ['year',key]]
    
    #  unpack prog
    jaar = prog[0]
    prognose_mean= prog[1]
    prognose_min= prog[2]
    prognose_max= prog[3]
    
    prognose = realisatie[:][-2:-1]
    
    prognose = prognose.append({'year': jaar,
                                key: prognose_mean}, 
                                ignore_index=True)
    prognose['max'] = prognose[key]
    prognose['max'][-1:] = prognose_max
    prognose['max'] = prognose['max'] - prognose[key]
    prognose['min'] = prognose[key]
    prognose['min'][-1:] = prognose_min
    prognose['min'] = prognose[key] - prognose['min']
    
    # set ylim
    if ylim is None:
        ylim = [0,max([prognose_max,max(realisatie[key])])+1000]   
    
    #%% plot
    
    fig1, ax1 = plt.subplots(figsize=(12, 4))
    
    # plot realisatie
    plt.plot(realisatie['year'],
             realisatie[key],
             marker='o',
             markersize=markersize,
             markeredgewidth=2,
             fillstyle='none',
             color=c1,
             linewidth=linesize,
             label = 'realisatie')
    
    # plot prognose
    plt.errorbar(prognose['year'],
             prognose[key],
             yerr= [prognose['min'],prognose['max']],
             elinewidth =linesize,
             capsize =4,
             ecolor= c3,
             marker='_',
             markersize=markersize,
             markeredgewidth=4,
             markeredgecolor=c3,
             fillstyle='none',
             color=c2,
             linewidth=linesize,
             label = 'prognose')
    
    
#    plt.fill_between(prognose['year']-1, 
#                     prognose[key]-prognose['min'], 
#                     prognose[key]+prognose['max'], 
#                     color=c4, 
#                     alpha=0.3)
#    
    # set yticks format
    plt.yticks(fontname = font,
               fontsize = fontsize)
    
    # set xticks format
    plt.xticks(np.arange(min(realisatie['year']), 
                         jaar+1, 
                         1.0),
               fontname = font,
               fontsize = fontsize)
    ax1 = plt.gca()
    ax1.set_ylim(ylim)
    ax1.grid(axis='y')
    
    # get rid of box around grpah
    for spine in plt.gca().spines.values():
        spine.set_visible(False) #Indentation updated..
        
    # Add a legend
    plt.legend(prop={'family':font, 'size':fontsize},
               ncol=2,
               bbox_to_anchor=(0.9, 1.15))
    
    # save
    plt.savefig(fn,dpi=300, bbox_inches='tight')
    
    # Show the plot
    plt.show()
    
def compGWCforGP(hdr_den,hdr_n,dat_den,dat_n):
    # compute GWC
    # lees woningbestand
    fn = 'lib/wbs2005.h5'
    wbs = ld.read_adressen(fn)
    w58den = []
    egh48den = []
    w48n = []
    sv40n = []
    
    for dat in dat_den:
        # interpoleer op de adressenlijst
        db_den, func_den = ld.grid_interpolatie(wbs, hdr_den, dat)
        # tellen
        w_den, egh  = ld.tellen_etmaal(db_den, wbs['woningen'], wbs['personen']) 
        print('aantal woningen binnen 48 Lden = '+ str(w_den[1]))
        w58den.append(w_den[0])
        egh48den.append(egh[1])
    
    for dat in dat_n:
        db_n, func_n = ld.grid_interpolatie(wbs, hdr_n, dat)
        w_n, sv     = ld.tellen_nacht(db_n,  wbs['woningen'], wbs['personen'])
        w48n.append(w_n[0])
        sv40n.append(sv[1])        
    
    
    
    d = {'Criterium': ['Won 58 dB(A) Lden',
                       'EGH 48 dB(A) Lden',
                       'Won 48 dB(A) Lnight',
                       'SV 40 dB(A) Lnight'],
         'mean': [sum(w58den)/len(w58den), 
                  sum(egh48den)/len(egh48den),
                  sum(w48n)/len(w48n), 
                  sum(sv40n)/len(sv40n)],
         'max': [max(w58den), max(egh48den), max(w48n), max(sv40n)],
         'min': [min(w58den), min(egh48den), min(w48n), min(sv40n)]}
    
    GWC = pd.DataFrame(data=d)
    GWC['mean'] = GWC['mean'].round(-2)
    GWC['max'] = GWC['max'].round(-2)
    GWC['min'] = GWC['min'].round(-2)
    
    return GWC

def fig23(DEN_zw1,label1,fn,hs,
                   DEN_zw2=None,
                   label2 =None,
                   DEN_zw3=None,
                   label3 =None):
    c1 = hs['kleuren']['schipholblauw']
    c2 = hs['kleuren']['middagblauw']
    c3 = hs['kleuren']['wolkengrijs_1']
    c4 = hs['kleuren']['schemergroen']
   
    font = hs['fontname']['grafiek']
    fontsize = hs['size']['font']
   
    r =range(len(DEN_zw1))
    names =DEN_zw1.index.tolist()
   
    fig1, ax1 = plt.subplots(figsize=(12, 4))
    ax1.grid(axis='y')
  
    
    ax1.bar([x-0.2 for x in r],
         DEN_zw1,
         align='center',
         color=c1,
         width= 0.2,
         label=label1)
    
    if DEN_zw2 is not None:
        ax1.bar(r,
             DEN_zw2,
             align='center',
             color=c2,
             width= 0.2,
             label=label2)
    if DEN_zw3 is not None:
        ax1.bar([x+0.2 for x in r],
             DEN_zw3,
             align='center',
             color=c3,
             width= 0.2,
             label=label3)
 
    plt.xticks(r,
               names,             
               fontname = font,
               fontsize = fontsize)
   
    vals = ax1.get_yticks()
    ax1.set_yticklabels(['{0:.0f}%'.format(x) for x in vals],             
               fontname = font,
               fontsize = fontsize)
    
    plt.xlabel('Maximum startgewicht in tonnen',              
               fontname = font,
               fontsize = fontsize)
   
    plt.legend()
    # get rid of box around grpah
    for spine in plt.gca().spines.values():
        spine.set_visible(False) #Indentation updated..
 
    # save
    plt.savefig(fn,dpi=300, bbox_inches='tight')
   
    plt.show()
    
def procedureverdeling(realisatie,prognose):
    
    #%% tabel 2.3

    # realisatie
    data_realisatie_HV_D = realisatie.loc[realisatie['C_AD']=='realisatie, starts']
    
    data_realisatie_HV_D['Procedure'] = data_realisatie_HV_D['C_Klasse']
    data_realisatie_HV_D['Procedure'] = 'NADP1'
    data_realisatie_HV_D.loc[data_realisatie_HV_D['C_Klasse']>=600,'Procedure'] = 'NADP2'
    
    realisatie_starts_verdeling = data_realisatie_HV_D.groupby(['Procedure'])['Procedure'].count()
    # format
    realisatie_starts_verdeling = round(realisatie_starts_verdeling/realisatie_starts_verdeling.sum()*100,1)
    
    
    # prognose
    data_prognose_mean_D = prognose.loc[prognose['d_lt']=='T']
    
    data_prognose_mean_D['Procedure'] = data_prognose_mean_D['d_proc']
    data_prognose_mean_D['Procedure'] = 'NADP1'
    data_prognose_mean_D.loc[data_prognose_mean_D['d_proc']>=600,'Procedure'] = 'NADP2'
    
    prognose_starts_verdeling = data_prognose_mean_D.groupby(['Procedure'])['total'].sum()
    
    # format
    prognose_starts_verdeling = round(prognose_starts_verdeling/prognose_starts_verdeling.sum()*100,1)
    
    # tabel maken
    d = {'prognose': prognose_starts_verdeling, 'realisatie': realisatie_starts_verdeling}
    starts = pd.DataFrame(data=d)
    del starts.index.name
#    print(starts)
    
    
    #%% tabel 2.4
    
    # realisatie
    data_realisatie_HV_A = realisatie.loc[realisatie['C_AD']=='realisatie, landingen']
    
    data_realisatie_HV_A['Procedure'] = data_realisatie_HV_A['C_Klasse'].fillna(0).astype(int) % 10
    
    data_realisatie_HV_A.loc[data_realisatie_HV_A['Procedure']==0,'Procedure'] = '2000 [ft]'
    data_realisatie_HV_A.loc[data_realisatie_HV_A['Procedure']==1,'Procedure'] = '3000 [ft]'
    data_realisatie_HV_A.loc[data_realisatie_HV_A['Procedure']==9,'Procedure'] = 'CDA'
    
    realisatie_landingen_verdeling = data_realisatie_HV_A.groupby(['Procedure'])['Procedure'].count()
    # format
    realisatie_landingen_verdeling = round(realisatie_landingen_verdeling/realisatie_landingen_verdeling.sum()*100,1)
    
    
    # prognose
    data_prognose_mean_A = prognose.loc[prognose['d_lt']=='L']
    
    data_prognose_mean_A['Procedure'] = data_prognose_mean_A['d_proc'].fillna(0).astype(int) % 10
    
    data_prognose_mean_A.loc[data_prognose_mean_A['Procedure']==0,'Procedure'] = '2000 [ft]'
    data_prognose_mean_A.loc[data_prognose_mean_A['Procedure']==1,'Procedure'] = '3000 [ft]'
    data_prognose_mean_A.loc[data_prognose_mean_A['Procedure']==9,'Procedure'] = 'CDA'
    
    prognose_landingen_verdeling = data_prognose_mean_A.groupby(['Procedure'])['total'].sum()
    
    # format
    prognose_landingen_verdeling = round(prognose_landingen_verdeling/prognose_landingen_verdeling.sum()*100,1)
    
    # tabel maken
    d = {'prognose': prognose_landingen_verdeling, 'realisatie': realisatie_landingen_verdeling}
    landingen = pd.DataFrame(data=d)
    del landingen.index.name
#    print(landingen)
    
    return starts, landingen