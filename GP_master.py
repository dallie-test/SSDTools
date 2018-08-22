import numpy as np
import pandas as pd
#pd.options.display.float_format = '{:.0f}'.format
import lib.GPlib as lg
import lib.doc29lib as ld

import matplotlib.pyplot as plt
from matplotlib import rc

import lib.huisstijl as lh
from matplotlib.ticker import FormatStrFormatter
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from scipy.misc import imread
import matplotlib.cbook as cbook
import lib.plotformat as pformat
from matplotlib import rcParams
import os

#%% inputs folders
input_folder1       = '../4.b Output Daisy/Hybride windroos met groot onderhoud/'
input_folder2       = '../4.b Output Daisy/Hybride windroos/'
input_folder3       = '../4.b Output Daisy/Hybride windroos met groot onderhoud incl GA grid/'
winter_scenarios    = ['Winter onverstoord',
                       'GOH 18C36C winter']
zomer_scenarios     = ['Zomer onverstoord',
                       'Zomer onverstoord deel 2',
                       'GOH 0624',
                       'GOH 18C36C zomer']
#winter_scenarios    = ['Winterseizoen']
#zomer_scenarios     = ['Zomerseizoen']

output_folder       = 'output/'


#%% process inputs

# check if output folder exists:
if not os.path.exists(output_folder):
    os.makedirs(output_folder)

if not os.path.exists(output_folder+ 'figuren/'):
    os.makedirs(output_folder+ 'figuren/')   
    
# read winter and zomer traffics
winter_scenarios    = [input_folder1 + 'traffic '+ s + '.txt' for s in winter_scenarios]
zomer_scenarios     = [input_folder1 + 'traffic '+ s + '.txt' for s in zomer_scenarios]

# read winter and summer traffics
tf_winter = lg.appendDF(winter_scenarios)
tf_zomer = lg.appendDF(zomer_scenarios)

# read total trafic
tf = pd.read_csv(input_folder1 + 'traffic 1971-2017 - mean.txt', delim_whitespace = 1)

# read preference traffic
tf_pref = pd.read_csv(input_folder1 + 'traffic 1971-2017 - pref.txt', delim_whitespace = 1)

# baancombinaties 
baancombinaties= pd.read_csv('input/Baancombinaties.txt', delim_whitespace = 1)

# routesector
routesector = pd.read_csv('input/routesector.txt', delim_whitespace = 1)

# huisstijl
hs = lh.getHuisStijl() 
rcParams['font.family'] = hs['fontname']['grafiek']

# template
template = 'input/gebruiksprognose_template.docx'
doc = Document(template)
styles = doc.styles
table_header = styles['Stijl2']
table_body_big =  styles['Stijl3']
table_body_small =  styles['Stijl4']
tables = doc.tables
paragraphs = doc.paragraphs

# history excel
history = pd.read_excel('input/history.xlsx')

#%% FIGUUR 2.1

# prognose data 
jaar = 2019
prognose_max = 500000
prog = sum(tf['total'])
prognose_min = 492000
prognose = [jaar,prog,prognose_min,prognose_max]

# make figure
fn = output_folder +'figuren/figuur21.png'
lg.figHistory(history,prognose,'verkeer',fn,hs,ylim = [350000,505000])

#% input figure into template
p = tables[0].rows[0].cells[0].paragraphs[0]
p.add_run().add_picture(fn,
              width=Inches(6.5), 
              height=Inches(2.3))
p.alignment = WD_ALIGN_PARAGRAPH.CENTER


#%% TABEL 2.2
tf.loc[tf['d_schedule']==6,'d_den']='vroege ochtend'
tf.loc[tf['d_den']=='N','d_den']='nacht'
tf.loc[tf['d_den']=='D','d_den']='dag'
tf.loc[tf['d_den']=='E','d_den']='avond'
tf.loc[tf['d_lt']=='T','d_lt']='starts'
tf.loc[tf['d_lt']=='L','d_lt']='landingen'    

DEN = pd.pivot_table(tf,index ='d_den',columns='d_lt',values='total',aggfunc=np.sum)
DEN['totaal']=DEN['starts']+DEN['landingen']
DEN.loc['totaal'] = pd.Series(DEN['totaal'].sum(), index = ['totaal'])
del DEN.index.name
DEN.index.name = 'periode'
DEN = DEN.reset_index()
DEN = DEN.fillna('')

print(DEN)

lg.df2table(tables[1],DEN,table_header,table_body_big,table_body_small)

#%% FIGUUR 2.2
DEN_winter = lg.DENverdeling(tf_winter,'d_schedule','total','d_lt','winter')
DEN_zomer = lg.DENverdeling(tf_zomer,'d_schedule','total','d_lt','zomer')
DEN_zw = pd.merge(DEN_zomer,DEN_winter,left_index=True,right_index=True,how='left')
DEN_zw = DEN_zw.transpose()
print(DEN_zw)

# make figure
fn = output_folder+ 'figuren/figuur22.png'
lg.fig22(DEN_zw,fn,hs)

#% input figure into template
p = tables[2].rows[0].cells[0].paragraphs[0]
p.add_run().add_picture(fn,
              width=Inches(6.5), 
              height=Inches(2.3))
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

#%% FIGUUR 2.3
vvc_pattern     = ['[0]\/[0-9]','[12]\/[0-9]','[3]\/[0-9]','[45]\/[0-9]','[6]\/[0-9]','[7]\/[0-9]','[89]\/[0-9]']
MTOW            = ['< 6','6 - 40','40 - 60','60 - 160','160 - 230','230 - 300','> 300'] 

for find,replace in zip(vvc_pattern,MTOW): 
    tf = tf.replace(to_replace=find,value=replace,regex=True)

vloot = tf.groupby(['d_ac_cat'])['total'].sum()
vloot = vloot.reindex(MTOW)
vloot= vloot/vloot.sum()*100
vloot = vloot.fillna(0)
print(vloot)

# make figure
fn = output_folder+ 'figuren/figuur23.png'
lg.fig23(vloot,fn,hs)

#% input figure into template
p = tables[3].rows[0].cells[0].paragraphs[0]
p.add_run().add_picture(fn,
              width=Inches(6.5), 
              height=Inches(2.3))
p.alignment = WD_ALIGN_PARAGRAPH.CENTER


#%% FIGUUR 2.4
# only do caculations to check GPtools
tf = pd.merge(tf,routesector,left_on='d_route',right_on='route',how='left')
sectorverdeling = tf.groupby(['sector'])['total'].sum()
SIDverdeling= sectorverdeling[0:4]/sectorverdeling[0:4].sum()*100
STARverdeling = sectorverdeling[5:]/sectorverdeling[5:].sum()*100
print(SIDverdeling)
print(STARverdeling)

# TO DO: put graph figure 2.4 here

#%% TABEL 3.1
d = {'Aspect': ['Totaal Aantal vliegbewegingen',
                'Aantal vliegbewegingen in de nacht'],
    'Grens': ['Tot en met 2020 maximaal 500.000 vliegtuigbewegingen handelsverkeer op jaarbasis',
              'Maximaal 32.000 vliegtuigbewegingen tussen 23:00 uur en 07:00 uur'],
    'Prognose 2019': [float(DEN['totaal'][-1:]),
                      DEN['totaal'][2:4].sum()]}

tabel31 = pd.DataFrame(data=d)
lg.df2table(tables[4],tabel31,table_header,table_body_big,table_body_small)

#%% TABEL 4.3 


tabel42a, tabel42b = lg.tab42(tf_pref,baancombinaties)

    
print(tabel42a)
print(tabel42b)
# print to word 
lg.df2table(tables[5],tabel42a,table_header,table_body_big,table_body_small)

# print to word 
lg.df2table(tables[6],tabel42b,table_header,table_body_big,table_body_small)

#%% FIGUUR 4.3
style='MER'
labels = ['GP2019','GP2019 excl. GO']
trf_files = [input_folder1 + 'traffic 1971-2017 - years.txt',
             input_folder2 + 'traffic 1971-2017 - years.txt']
fn = 'output/figuren/figuur43.png'
den=['D', 'E', 'N']
lg.plot_baangebruik(trf_files,
                 labels,
                 hs,
                 den,
                 fn)

#% input figure into template
p = tables[7].rows[0].cells[0].paragraphs[0]
p.add_run().add_picture(fn,
              width=Inches(6.5), 
              height=Inches(2.3))
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

#% FIGUUR 4.4
fn = output_folder+ 'figuren/figuur44.png'
den=['N']
lg.plot_baangebruik(trf_files,
                 labels,
                 hs,
                 den,
                 fn,
                 ylim=[0,12000],
                 dy=1000)

#% input figure into template
p = tables[8].rows[0].cells[0].paragraphs[0]
p.add_run().add_picture(fn,
              width=Inches(6.5), 
              height=Inches(2.3))
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

#%% TABEL 4.4
lts = ['landingen','starts','landingen','starts']
dens = ['den','den','n','n']
keys = ['Aantal landingen','Aantal starts','Aantal landingen','Aantal starts']
ii  = [9,10,11,12]
ns  = [7,7,4,4]

for lt,key,den,i,n in zip(lts,keys,dens,ii,ns):
    
    # groupby runway
    if den=='n':
        baangebruik =  tf[(tf['d_lt']==lt) & ((tf['d_den']=='nacht') | (tf['d_den']=='vroege ochtend'))].groupby(['d_runway'])['total'].sum()
    else:
        baangebruik =  tf[tf['d_lt']==lt].groupby(['d_runway'])['total'].sum()
    
    # make nice dataframe
    baangebruik = baangebruik.to_frame().reset_index()
    baangebruik.columns = ['Baan', key]
    baangebruik = baangebruik.sort_values(key,ascending=False)
    tabel = baangebruik.nlargest(n,key)
    
    tabel = tabel.append({'Baan': 'overig',
                                key:baangebruik[key].sum()-tabel[key].sum()}, 
                                ignore_index=True)
    
    print(tabel)
    lg.df2table(tables[i],tabel,table_header,table_body_big,table_body_small)
    

#%% FIGUUR 5.1
hdr_den, dat_den = ld.gridimport(grid_dir = input_folder3,
                                           noise = 'GP2019 - Lden v01',
                                           scale = 1.0) 

# initieer plot
fig,ax = ld.start_contourplot()

# plot contouren
cs1 = lg.plotContour(ax,hdr_den,dat_den,48,
               hs['kleuren']['schipholblauw'],
               hs['kleuren']['middagblauw'])

cs2 = lg.plotContour(ax,hdr_den,dat_den,58,
               hs['kleuren']['schemergroen'],
               hs['kleuren']['wolkengrijs_1'])

# Legenda
cs1.collections[0].set_label('48 dB(A)')
cs2.collections[0].set_label('58 dB(A)')

legend = ax.legend(title=r'Etmaalperiode $L_{den}$',
                   fontsize = hs['size']['font'],
                   frameon=False,
                   loc='upper left',
                   bbox_to_anchor=(0.02, 0.96))

plt.setp(legend.get_title(),
         fontsize=hs['size']['font'])

fn = output_folder+ 'figuren/figuur51.png'
fig.savefig(fn, dpi=300)   
plt.close(fig)

#% input figure into template
p = tables[13].rows[0].cells[0].paragraphs[0]
p.add_run().add_picture(fn,
              width=Inches(6.5), 
              height=Inches(2.3))
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

#%% FIGUUR 5.2
hdr_n, dat_n = ld.gridimport(grid_dir = input_folder3,
                                           noise = 'GP2019 - Lnight v01',
                                           scale = 1.0) 

# initieer plot
fig,ax = ld.start_contourplot()

# plot contouren
cs1 = lg.plotContour(ax,hdr_n,dat_n,40,
               hs['kleuren']['schipholblauw'],
               hs['kleuren']['middagblauw'])

cs2 = lg.plotContour(ax,hdr_n,dat_n,48,
               hs['kleuren']['schemergroen'],
               hs['kleuren']['wolkengrijs_1'])

# Legenda
cs1.collections[0].set_label('40 dB(A)')
cs2.collections[0].set_label('48 dB(A)')

legend = ax.legend(title=r'Nachtperiode $L_{night}$',
                   fontsize = hs['size']['font'],
                   frameon=False,
                   loc='upper left',
                   bbox_to_anchor=(0.02, 0.96))

plt.setp(legend.get_title(),
         fontsize=hs['size']['font'])

fn = output_folder+ 'figuren/figuur52.png'
fig.savefig(fn, dpi=300)   
plt.close(fig)

#% input figure into template
p = tables[14].rows[0].cells[0].paragraphs[0]
p.add_run().add_picture(fn,
              width=Inches(6.5), 
              height=Inches(2.3))
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

#%% FIGUUR 5.3 - 5.6
key = ['w58den', 'eh48den','w48n', 'sv40n']
files = ['figuur53', 'figuur54', 'figuur55', 'figuur56']
fn = [output_folder + 'figuren/'+ f + '.png' for f in files]
jaar = 2019

GWC = lg.compGWCforGP(hdr_den,hdr_n,dat_den['dat'],dat_n['dat'])


i =0
for k,f in zip(key,fn):
    # prognose data 
    prognose = [jaar,
                GWC['mean'][i],
                GWC['max'][i],
                GWC['min'][i]]
    
    lg.figHistory(history,prognose,k,f,hs)
    #% input figure into template
    p = tables[15+i].rows[0].cells[0].paragraphs[0]
    p.add_run().add_picture(f,
                  width=Inches(6.5), 
                  height=Inches(2.3))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    i += 1

#%% save that bitch
doc.save(output_folder+ 'gebruiksprognose_2019_concept.docx')
